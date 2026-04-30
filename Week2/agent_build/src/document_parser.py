"""
Document parser for inbound vendor contracts.
Supports .docx format only (per CLAUDE.md §6.3).
Derived from D4 T-03 and D5 §2 (Word parsing).
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

try:
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@dataclass
class DocumentSection:
    heading: str
    body: str
    heading_level: int  # 1 = top-level heading; 0 = inferred heading (bold/caps)


@dataclass
class ParsedDocument:
    filename: str
    full_text: str
    sections: list[DocumentSection]
    page_count_estimate: int
    raw_headings: list[str]
    parse_error: Optional[str] = None


def parse_docx(filepath: str | Path) -> ParsedDocument:
    """
    Parses a .docx file into structured text with section boundaries.
    Returns a ParsedDocument; on failure, sets parse_error and returns empty structure.
    """
    filepath = Path(filepath)

    if not filepath.exists():
        return ParsedDocument(
            filename=filepath.name,
            full_text="",
            sections=[],
            page_count_estimate=0,
            raw_headings=[],
            parse_error=f"File not found: {filepath}",
        )

    if not DOCX_AVAILABLE:
        return ParsedDocument(
            filename=filepath.name,
            full_text="",
            sections=[],
            page_count_estimate=0,
            raw_headings=[],
            parse_error=(
                "python-docx is not installed. "
                "Run: pip install python-docx"
            ),
        )

    try:
        doc = DocxDocument(str(filepath))
    except Exception as exc:
        return ParsedDocument(
            filename=filepath.name,
            full_text="",
            sections=[],
            page_count_estimate=0,
            raw_headings=[],
            parse_error=f"Document parse error: {exc}",
        )

    sections: list[DocumentSection] = []
    current_heading: Optional[str] = None
    current_level: int = 0
    current_body: list[str] = []
    raw_headings: list[str] = []
    all_text_parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        all_text_parts.append(text)
        heading_level = _heading_level(para, text)

        if heading_level > 0:
            if current_heading is not None:
                sections.append(DocumentSection(
                    heading=current_heading,
                    body="\n".join(current_body),
                    heading_level=current_level,
                ))
            current_heading = text
            current_level = heading_level
            current_body = []
            raw_headings.append(text)
        else:
            current_body.append(text)

    if current_heading is not None:
        sections.append(DocumentSection(
            heading=current_heading,
            body="\n".join(current_body),
            heading_level=current_level,
        ))

    full_text = "\n".join(all_text_parts)
    return ParsedDocument(
        filename=filepath.name,
        full_text=full_text,
        sections=sections,
        page_count_estimate=_estimate_page_count(full_text),
        raw_headings=raw_headings,
    )


def _heading_level(para, text: str) -> int:
    """
    Returns 1–3 for headings, 0 for body paragraphs.
    Checks Word style name first, then heuristics.
    """
    style_name = para.style.name if para.style else ""

    if style_name.startswith("Heading "):
        try:
            return int(style_name.split(" ")[1])
        except (IndexError, ValueError):
            return 1

    if style_name in ("Title", "Subtitle"):
        return 1

    # Heuristics: short ALL CAPS, short bold, or numbered section header
    if len(text) > 120:
        return 0

    if text.isupper() and 1 <= len(text.split()) <= 12:
        return 1

    if re.match(r"^\d+(\.\d+)*\.?\s+[A-Z]", text) and len(text) <= 80:
        return 2

    # All-bold short paragraph
    if para.runs and all(run.bold for run in para.runs if run.text.strip()):
        if len(text.split()) <= 15:
            return 2

    return 0


def _estimate_page_count(text: str) -> int:
    """Estimates pages from word count (~500 words/page for dense legal text)."""
    word_count = len(text.split())
    return max(1, round(word_count / 500))
