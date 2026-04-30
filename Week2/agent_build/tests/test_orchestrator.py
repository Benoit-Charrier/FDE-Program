"""
Integration tests for the Orchestrator using mock dependencies and StubClauseClassifier.
No API keys or external services required.
"""
import pytest
from datetime import datetime
from pathlib import Path

from src.classifier import StubClauseClassifier, ClassificationResult
from src.config import PLAYBOOK_PATH
from src.hitl_queue import MockHITLQueue
from src.ironclad_client import MockIroncladClient
from src.models import (
    Contract,
    ContractStatus,
    PlaybookMatchStatus,
    RoutingClassification,
    TaskUnitType,
)
from src.orchestrator import Orchestrator
from src.playbook import PlaybookLoader


def make_contract(vendor="VendorCo Ltd", email="p@v.com") -> Contract:
    return Contract(
        vendor_name=vendor,
        vendor_email=email,
        date_received=datetime(2024, 4, 1, 9, 0),
        document_filename="test_contract.docx",
    )


def make_orchestrator(classifier=None) -> tuple[Orchestrator, MockIroncladClient, MockHITLQueue]:
    ironclad = MockIroncladClient()
    hitl = MockHITLQueue()
    playbook = PlaybookLoader(playbook_path=PLAYBOOK_PATH)
    orch = Orchestrator(
        classifier=classifier or StubClauseClassifier(),
        playbook=playbook,
        ironclad=ironclad,
        hitl_queue=hitl,
    )
    return orch, ironclad, hitl


# We need a real .docx file to test parsing; we create a minimal one programmatically.
@pytest.fixture
def minimal_docx(tmp_path: Path) -> Path:
    """Creates a test .docx with sections for all 7 clause types and enough words (~3,000) to pass the page threshold."""
    try:
        from docx import Document

        # ~430-word filler block to pad each section above the 500-word/page estimate
        _FILLER = (
            "This agreement sets out the terms and conditions under which the parties agree to "
            "conduct business with one another. Both parties acknowledge that this clause has been "
            "reviewed, negotiated, and agreed upon by legal counsel. The provisions herein are "
            "intended to be commercially reasonable and consistent with standard market practice "
            "for agreements of this type in the United Kingdom. In the event of any inconsistency "
            "between this clause and other provisions of the agreement, the more specific provision "
            "shall prevail unless otherwise agreed in writing by both parties. The parties further "
            "acknowledge that any amendment to this clause must be made in writing and signed by "
            "authorised representatives of both parties. This clause shall survive the termination "
            "or expiry of the agreement to the extent necessary to give effect to its provisions. "
        ) * 5  # repeat to reach ~400 words per section

        doc = Document()
        doc.add_heading("1. Limitation of Liability", level=1)
        doc.add_paragraph(
            "Vendor liability is capped at £300,000 or 12 months fees, whichever is greater. "
            + _FILLER
        )
        doc.add_heading("2. Data Processing Agreement", level=1)
        doc.add_paragraph(
            "The parties agree to process personal data in accordance with UK GDPR and DPA 2018. "
            "Sub-processor list is provided. Data residency is UK/EEA. Breach notification is 72 hours. "
            + _FILLER
        )
        doc.add_heading("3. Termination", level=1)
        doc.add_paragraph(
            "Either party may terminate with 30 days written notice for convenience. "
            + _FILLER
        )
        doc.add_heading("4. Intellectual Property Ownership", level=1)
        doc.add_paragraph(
            "All IP created by the Vendor for Helix under this agreement is owned by Helix. "
            "Background IP remains with the originating party. "
            + _FILLER
        )
        doc.add_heading("5. Service Level Agreement", level=1)
        doc.add_paragraph(
            "Vendor guarantees 99.5% uptime with a 4-hour critical issue response time. "
            "Service credits apply at 10% per 0.5% downtime below SLA. "
            + _FILLER
        )
        doc.add_heading("6. Governing Law", level=1)
        doc.add_paragraph(
            "This agreement is governed by English law. "
            "English courts have exclusive jurisdiction over disputes. "
            + _FILLER
        )
        doc.add_heading("7. Indemnification", level=1)
        doc.add_paragraph(
            "Vendor indemnifies Helix against third-party IP infringement claims arising from "
            "vendor deliverables and vendor gross negligence. "
            + _FILLER
        )
        path = tmp_path / "test_contract.docx"
        doc.save(str(path))
        return path
    except ImportError:
        pytest.skip("python-docx not installed — skipping orchestrator integration tests")


class TestOrchestratorStandardPath:
    def test_all_compliant_produces_standard_routing(self, minimal_docx):
        orch, ironclad, hitl = make_orchestrator()
        contract = make_contract()
        run = orch.process_contract(contract, minimal_docx)

        assert run.completed, f"Run failed: {run.error}"
        assert run.routing_classification == RoutingClassification.STANDARD
        # hitl_required is True because the fixture includes a DPA clause, which always
        # triggers ET-2 unconditionally — routing classification and hitl_required are independent.
        # Status is AWAITING_APPROVAL (not REVIEWED_STANDARD) because HITL path was taken.
        assert run.review_decision is not None
        assert run.review_decision.approval_token is None  # agent never sets this

    def test_ironclad_case_created(self, minimal_docx):
        orch, ironclad, hitl = make_orchestrator()
        contract = make_contract()
        run = orch.process_contract(contract, minimal_docx)

        assert run.ironclad_case_id.startswith("IRONCLAD-MOCK-")
        case = ironclad.get_case(contract.contract_id)
        assert case is not None
        assert len(case.clause_reviews) == 7

    def test_clause_reviews_written_for_all_7_types(self, minimal_docx):
        orch, ironclad, _ = make_orchestrator()
        run = orch.process_contract(make_contract(), minimal_docx)

        assert len(run.clause_reviews) == 7
        found_types = {r.task_unit_type for r in run.clause_reviews}
        assert found_types == set(TaskUnitType)

    def test_dpa_clause_always_triggers_hitl(self, minimal_docx):
        """ET-2: DPA clause is always flagged regardless of confidence."""
        orch, _, hitl = make_orchestrator()
        run = orch.process_contract(make_contract(), minimal_docx)

        assert run.hitl_required  # DPA clause in fixture → always HITL
        et2_payloads = [p for p in hitl.all_items() if p.trigger_id == "ET-2"]
        assert len(et2_payloads) == 1


class TestOrchestratorHITLPath:
    def test_major_deviation_triggers_hitl_and_awaiting_approval(self, minimal_docx):
        classifier = StubClauseClassifier(
            results={
                TaskUnitType.IP_OWNERSHIP: ClassificationResult(
                    task_unit_type=TaskUnitType.IP_OWNERSHIP,
                    playbook_match_status=PlaybookMatchStatus.MAJOR_DEVIATION,
                    agent_confidence_score=0.88,
                    agent_reasoning_summary="Vendor claims ownership of all work product.",
                    playbook_section_retrieved="IP_OWNERSHIP section",
                )
            }
        )
        orch, _, hitl = make_orchestrator(classifier)
        run = orch.process_contract(make_contract(), minimal_docx)

        assert run.hitl_required
        assert run.contract.status == ContractStatus.AWAITING_APPROVAL
        assert run.routing_classification == RoutingClassification.ESCALATION_REQUIRED

        et4_payloads = [p for p in hitl.all_items() if p.trigger_id == "ET-4"]
        assert len(et4_payloads) >= 1

    def test_low_confidence_triggers_et1_hitl(self, minimal_docx):
        classifier = StubClauseClassifier(default_confidence=0.70)
        orch, _, hitl = make_orchestrator(classifier)
        run = orch.process_contract(make_contract(), minimal_docx)

        assert run.hitl_required
        et1_payloads = [p for p in hitl.all_items() if p.trigger_id == "ET-1"]
        assert len(et1_payloads) > 0


class TestOrchestratorHardStops:
    def test_missing_docx_file_halts_run(self):
        orch, _, _ = make_orchestrator()
        run = orch.process_contract(make_contract(), Path("/nonexistent/contract.docx"))
        assert not run.completed
        assert run.error is not None

    def test_non_docx_filename_rejected_at_model_creation(self):
        with pytest.raises(Exception):
            Contract(
                vendor_name="V",
                vendor_email="v@v.com",
                date_received=datetime.utcnow(),
                document_filename="contract.pdf",  # must fail validation
            )

    def test_approval_token_always_none_from_agent(self, minimal_docx):
        orch, _, _ = make_orchestrator()
        run = orch.process_contract(make_contract(), minimal_docx)
        if run.review_decision:
            assert run.review_decision.approval_token is None


class TestVendorHistory:
    def test_prior_escalation_triggers_et6(self, minimal_docx):
        orch, ironclad, hitl = make_orchestrator()
        # Simulate a prior escalation-required case for the same vendor
        from datetime import datetime
        prior_contract = Contract(
            vendor_name="VendorCo Ltd",
            vendor_email="p@v.com",
            date_received=datetime(2024, 1, 1),
            document_filename="prior.docx",
        )
        ironclad.create_case(prior_contract)
        from uuid import uuid4
        from src.models import ReviewDecision, DecisionType
        ironclad.write_routing_decision(
            prior_contract.contract_id,
            ReviewDecision(
                contract_id=prior_contract.contract_id,
                clause_review_ids=[uuid4()],
                decision_type=DecisionType.ESCALATE,
                decision_made_by="AGENT",
                requires_lawyer_approval=False,
            ),
        )
        # The mock returns based on contract routing_classification field, not the decision.
        # Manually set the routing classification on the case record for the mock to detect.
        ironclad._cases[str(prior_contract.contract_id)].routing_classification = "ESCALATION_REQUIRED"

        run = orch.process_contract(make_contract(), minimal_docx)
        et6_payloads = [p for p in hitl.all_items() if p.trigger_id == "ET-6"]
        assert len(et6_payloads) == 1
