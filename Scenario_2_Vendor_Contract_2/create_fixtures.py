"""Run this once to generate sample DOCX fixtures in the 'Input Contracts' folder."""
from pathlib import Path
from docx import Document


def make_standard():
    doc = Document()
    doc.add_heading("Vendor Services Agreement — Standard", 0)
    doc.add_paragraph("This agreement is entered into between Acme Corp ('Company') and VendorCo Ltd ('Vendor').")

    doc.add_heading("1. Limitation of Liability", level=1)
    doc.add_paragraph(
        "Vendor's liability under this agreement shall be limited to the total fees paid "
        "by Company to Vendor in the twelve (12) months preceding the claim. In no event "
        "shall either party be liable for indirect or consequential damages."
    )

    doc.add_heading("2. Data Processing", level=1)
    doc.add_paragraph(
        "The parties shall execute a Data Processing Agreement attached hereto as Exhibit A, "
        "which is GDPR compliant and incorporates the standard contractual clauses as approved "
        "by the European Commission."
    )

    doc.add_heading("3. Termination", level=1)
    doc.add_paragraph(
        "Either party may terminate this agreement for convenience on 30 days written notice "
        "to the other party. Termination does not affect accrued rights or obligations."
    )

    doc.add_heading("4. Intellectual Property", level=1)
    doc.add_paragraph(
        "Each party retains ownership of pre-existing IP brought to this engagement. "
        "All deliverables created by Vendor for Company shall be work for hire owned by Company."
    )

    doc.add_heading("5. Service Level Agreement", level=1)
    doc.add_paragraph(
        "Vendor commits to 99.5% uptime for the platform, measured monthly. "
        "Scheduled maintenance windows are excluded from uptime calculations."
    )

    doc.add_heading("6. Governing Law", level=1)
    doc.add_paragraph(
        "This agreement is governed by the laws of England and Wales. "
        "The parties submit to the exclusive jurisdiction of the courts of England and Wales."
    )

    doc.add_heading("7. Indemnification", level=1)
    doc.add_paragraph(
        "Vendor shall indemnify Company against third-party IP infringement claims only, "
        "arising directly from Vendor's deliverables. This indemnity excludes claims arising "
        "from Company's modifications or combination with third-party materials."
    )

    Path("Input Contracts").mkdir(exist_ok=True)
    doc.save("Input Contracts/sample_standard.docx")
    print("[OK] Created Input Contracts/sample_standard.docx")


def make_negotiable():
    doc = Document()
    doc.add_heading("Vendor Services Agreement — Negotiable Deviations", 0)
    doc.add_paragraph("This agreement is entered into between Acme Corp ('Company') and DevCo Ltd ('Vendor').")

    doc.add_heading("1. Limitation of Liability", level=1)
    doc.add_paragraph(
        "Vendor's liability under this agreement shall be capped at 3x annual fees paid "
        "by Company in the contract year in which the claim arises. Neither party shall be "
        "liable for indirect or punitive damages."
    )

    doc.add_heading("2. Data Processing", level=1)
    doc.add_paragraph(
        "A Data Processing Agreement is attached hereto as Schedule B, which is GDPR compliant "
        "and incorporates the standard contractual clauses issued by the European Commission."
    )

    doc.add_heading("3. Termination", level=1)
    doc.add_paragraph(
        "Either party may terminate this agreement for convenience on 60 days written notice "
        "to the other party. All fees accrued prior to the termination date remain payable."
    )

    doc.add_heading("4. Intellectual Property", level=1)
    doc.add_paragraph(
        "Each party retains ownership of pre-existing IP. Deliverables created under this "
        "agreement shall be work for hire owned by Company upon full payment."
    )

    doc.add_heading("5. Service Levels", level=1)
    doc.add_paragraph(
        "Vendor shall use commercially reasonable efforts to maintain platform availability "
        "and shall provide monthly availability reports to Company."
    )

    doc.add_heading("6. Governing Law", level=1)
    doc.add_paragraph(
        "This agreement shall be governed by and construed in accordance with the laws of "
        "the State of New York, without regard to its conflict of law principles."
    )

    doc.add_heading("7. Indemnification", level=1)
    doc.add_paragraph(
        "Each party shall indemnify the other for mutual indemnification arising from "
        "gross negligence or wilful misconduct of the indemnifying party's personnel."
    )

    doc.save("Input Contracts/sample_negotiable.docx")
    print("[OK] Created Input Contracts/sample_negotiable.docx")


def make_escalation():
    doc = Document()
    doc.add_heading("Vendor Services Agreement — Escalation Required", 0)
    doc.add_paragraph("This agreement is entered into between Acme Corp ('Company') and RiskyCo Ltd ('Vendor').")

    doc.add_heading("1. Limitation of Liability", level=1)
    doc.add_paragraph(
        "Notwithstanding anything to the contrary, Vendor's liability under this agreement "
        "shall be unlimited. The parties agree that no cap shall apply to any claim of "
        "whatever nature arising under or in connection with this agreement."
    )

    doc.add_heading("2. Termination", level=1)
    doc.add_paragraph(
        "Either party may terminate this agreement on 30 days written notice. "
        "Early termination by Company shall trigger a penalty equal to remaining contract value."
    )

    doc.add_heading("3. Intellectual Property", level=1)
    doc.add_paragraph(
        "Each party retains ownership of pre-existing IP brought to this engagement. "
        "Deliverables shall be work for hire and owned by the Company upon delivery."
    )

    doc.add_heading("4. Service Levels", level=1)
    doc.add_paragraph(
        "Vendor commits to 99.9% monthly uptime for the core platform services as defined "
        "in Schedule A."
    )

    doc.add_heading("5. Governing Law", level=1)
    doc.add_paragraph(
        "This agreement shall be governed by the laws of England and Wales."
    )

    doc.add_heading("6. Indemnification", level=1)
    doc.add_paragraph(
        "Company shall indemnify for third-party IP infringement claims only arising "
        "directly from use of the Vendor platform."
    )

    # Note: DPA is intentionally absent to trigger ESC-07

    doc.save("Input Contracts/sample_escalation.docx")
    print("[OK] Created Input Contracts/sample_escalation.docx")


if __name__ == "__main__":
    make_standard()
    make_negotiable()
    make_escalation()
    print("\nAll fixtures created in 'Input Contracts/'.")
